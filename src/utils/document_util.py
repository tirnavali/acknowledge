"""Utilities for extracting text and metadata from Word documents."""

import os
import logging

logger = logging.getLogger(__name__)

DOCUMENT_EXTS = {".doc", ".docx"}


def is_document(file_path: str) -> bool:
    return os.path.splitext(file_path)[1].lower() in DOCUMENT_EXTS


def extract_docx_text(file_path: str) -> str | None:
    """Extract all paragraph text from a .docx file. Returns None for .doc or on error."""
    if os.path.splitext(file_path)[1].lower() != ".docx":
        return None
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs) if paragraphs else None
    except Exception as e:
        logger.warning(f"Could not extract text from {file_path}: {e}")
        return None


def extract_doc_metadata(file_path: str) -> dict:
    """Extract core properties from a .docx file (author, title, etc.)."""
    if os.path.splitext(file_path)[1].lower() != ".docx":
        return {}
    try:
        from docx import Document
        doc = Document(file_path)
        props = doc.core_properties
        meta = {}
        if props.author:
            meta["author"] = props.author
        if props.title:
            meta["title"] = props.title
        if props.subject:
            meta["subject"] = props.subject
        if props.created:
            meta["created"] = props.created.isoformat()
        if props.modified:
            meta["modified"] = props.modified.isoformat()
        para_count = len(doc.paragraphs)
        if para_count:
            meta["paragraph_count"] = para_count
        return meta
    except Exception as e:
        logger.warning(f"Could not extract metadata from {file_path}: {e}")
        return {}


def generate_document_thumbnail(dst_path: str, thumb_path: str) -> bool:
    """Generate a simple document-icon JPEG thumbnail for a doc/docx file."""
    try:
        from PIL import Image, ImageDraw, ImageFont
        import os

        img = Image.new("RGB", (300, 300), color=(245, 245, 250))
        draw = ImageDraw.Draw(img)

        # Page shape
        pw, ph = 160, 200
        px = (300 - pw) // 2
        py = (300 - ph) // 2
        corner = 18
        draw.rounded_rectangle([px, py, px + pw, py + ph], radius=corner, fill="white", outline="#b0b8c8", width=2)

        # Folded corner
        fold = 30
        draw.polygon([(px + pw - fold, py), (px + pw, py + fold), (px + pw - fold, py + fold)], fill="#dde3ee")
        draw.line([(px + pw - fold, py), (px + pw - fold, py + fold), (px + pw, py + fold)], fill="#b0b8c8", width=2)

        # Text lines (simulate content)
        line_color = "#c5cad8"
        lx0 = px + 18
        lx1 = px + pw - 18
        for i, ly in enumerate(range(py + 48, py + ph - 30, 16)):
            short = i % 5 == 4
            draw.line([(lx0, ly), (lx1 - (30 if short else 0), ly)], fill=line_color, width=3)

        # Extension badge
        ext = os.path.splitext(dst_path)[1].upper().lstrip(".")
        badge_x0, badge_y0 = px + 14, py + ph - 44
        badge_x1, badge_y1 = px + 14 + 50, py + ph - 44 + 22
        draw.rounded_rectangle([badge_x0, badge_y0, badge_x1, badge_y1], radius=5, fill="#4a6fa5")
        try:
            font = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 13)
        except Exception:
            font = ImageFont.load_default()
        draw.text(((badge_x0 + badge_x1) / 2, (badge_y0 + badge_y1) / 2),
                  ext, fill="white", font=font, anchor="mm")

        img.thumbnail((300, 300))
        img.save(thumb_path, "JPEG", quality=85)
        return True
    except Exception as e:
        logger.warning(f"Could not generate document thumbnail: {e}")
        return False
