"""
create_fake_test_archive.py — Creates a realistic fake archive with images, PDFs, Word docs, and videos
across multiple years to test the Batch Upload and Duplicate Prevention features.
"""

import os
import sys
from PIL import Image, ImageDraw, ImageFont
import numpy as np

# Ensure root is in path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))


def make_fake_image(path: str, title: str, color=(70, 130, 180)):
    """Creates a sample test JPEG image with a label."""
    os.makedirs(os.path.dirname(path), exist_ok=True)
    img = Image.new("RGB", (800, 600), color=color)
    draw = ImageDraw.Draw(img)

    # Decorative shapes
    draw.rectangle([40, 40, 760, 560], outline="white", width=4)
    draw.ellipse([300, 200, 500, 400], fill=(240, 240, 240), outline="black", width=2)

    try:
        font = ImageFont.load_default()
    except Exception:
        font = None

    draw.text((400, 450), title, fill="white", font=font, anchor="mm")
    img.save(path, "JPEG", quality=90)


def make_fake_docx(path: str, title: str, content: str):
    """Creates a sample .docx document with python-docx."""
    try:
        from docx import Document
        os.makedirs(os.path.dirname(path), exist_ok=True)
        doc = Document()
        doc.add_heading(title, 0)
        doc.add_paragraph(content)
        doc.add_paragraph("Toplantı gündem maddeleri başarıyla karara bağlanmıştır.")
        doc.core_properties.title = title
        doc.core_properties.author = "Acknowledge Test Generator"
        doc.save(path)
    except Exception as e:
        print(f"Warning: python-docx could not generate {path}: {e}")


def make_fake_pdf(path: str, title: str, text_content: str):
    """Creates a sample valid PDF file using PySide6's QPdfWriter or simple PDF format."""
    try:
        from PySide6 import QtGui, QtCore
        os.makedirs(os.path.dirname(path), exist_ok=True)
        writer = QtGui.QPdfWriter(path)
        writer.setPageSize(QtGui.QPageSize(QtGui.QPageSize.A4))
        painter = QtGui.QPainter(writer)
        
        font_title = QtGui.QFont("Arial", 16, QtGui.QFont.Bold)
        painter.setFont(font_title)
        painter.setPen(QtGui.QColor("#2b5b84"))
        painter.drawText(100, 200, title)

        font_body = QtGui.QFont("Arial", 11)
        painter.setFont(font_body)
        painter.setPen(QtGui.QColor("#333333"))
        painter.drawText(100, 400, text_content)
        painter.drawText(100, 550, "Bu belge Acknowledge Test Sistemi tarafından otomatik üretilmiştir.")
        painter.end()
    except Exception as e:
        print(f"Warning: QPdfWriter could not generate {path}: {e}")


def make_fake_mp4(path: str, duration_sec: int = 2):
    """Creates a small valid MP4 test video using OpenCV."""
    try:
        import cv2
        os.makedirs(os.path.dirname(path), exist_ok=True)
        fourcc = cv2.VideoWriter_fourcc(*'mp4v')
        fps = 10
        w, h = 320, 240
        out = cv2.VideoWriter(path, fourcc, fps, (w, h))

        for frame_idx in range(fps * duration_sec):
            frame = np.zeros((h, w, 3), dtype=np.uint8)
            frame[:, :] = (40 + frame_idx * 5, 80, 140)
            cv2.putText(frame, f"Frame {frame_idx}", (50, 120), cv2.FONT_HERSHEY_SIMPLEX, 0.8, (255, 255, 255), 2)
            out.write(frame)

        out.release()
    except Exception as e:
        print(f"Warning: OpenCV could not generate video {path}: {e}")


def create_test_archive(base_dir: str = "test_archive"):
    """
    Creates a comprehensive test directory structure with 2024 and 2025 events,
    including the duplicate test case ('01.01.2024 divan toplantısı' vs '01.01.2025 divan toplantısı').
    """
    abs_base = os.path.abspath(base_dir)
    print(f"🚀 Test arşivi oluşturuluyor: {abs_base}")

    events_structure = [
        # --- 2024 EVENTS ---
        ("2024 Yılı Etkinlikleri/01.01.2024 divan toplantısı", [
            ("baskan_konusma.jpg", "image", "2024 Divan Toplantısı Açılış"),
            ("yonetim_kurulu.png", "image", "2024 Divan Kurulu"),
            ("divan_tutanagi.pdf", "pdf", "2024 Yılı 1. Olağan Divan Tutanakları"),
            ("toplanti_kararlari.docx", "docx", "2024 Divan Toplantısı Kararları"),
            ("acilis_videosu.mp4", "video", "2024 Divan Video"),
        ]),
        ("2024 Yılı Etkinlikleri/15.03.2024 Gençlik Çalıştayı", [
            ("calistay_afis.jpg", "image", "Gençlik Çalıştayı 2024"),
            ("grup_calismasi.jpg", "image", "Grup Çalışmaları"),
            ("calistay_raporu.pdf", "pdf", "2024 Gençlik Çalıştayı Sonuç Bildirgesi"),
        ]),
        ("2024 Yılı Etkinlikleri/19.05.2024 Bayram Kutlaması", [
            ("toren_alani.jpg", "image", "19 Mayıs Tören Alanı"),
            ("odul_toreni.jpg", "image", "Ödül Töreni 2024"),
            ("bayram_programi.docx", "docx", "19 Mayıs Program Akışı"),
        ]),
        ("2024 Yılı Etkinlikleri/10.11.2024 Anma Töreni", [
            ("saygi_durusu.jpg", "image", "10 Kasım Saygı Duruşu"),
            ("anma_metni.pdf", "pdf", "10 Kasım Anma Konuşma Metni"),
        ]),

        # --- 2025 EVENTS (Includes duplicate named event!) ---
        ("2025 Yılı Etkinlikleri/01.01.2025 divan toplantısı", [
            ("yeni_baskan_konusma.jpg", "image", "2025 Divan Toplantısı Açılış"),
            ("yeni_katilimcilar.jpg", "image", "2025 Divan Katılımcıları"),
            ("yeni_divan_tutanagi.pdf", "pdf", "2025 Yılı 1. Olağan Divan Tutanakları"),
            ("yeni_kararlar.docx", "docx", "2025 Divan Toplantısı Yeni Kararlar"),
            ("yeni_acilis_videosu.mp4", "video", "2025 Divan Video"),
        ]),
        ("2025 Yılı Etkinlikleri/23.04.2025 Çocuk Şenliği", [
            ("senlik_oyunlar.jpg", "image", "Çocuk Şenliği Oyunlar"),
            ("senlik_programi.docx", "docx", "23 Nisan 2025 Programı"),
        ]),
        ("2025 Yılı Etkinlikleri/29.10.2025 Cumhuriyet Balosu", [
            ("balo_acilis.jpg", "image", "Cumhuriyet Balosu 2025"),
            ("dans_gosterisi.jpg", "image", "Cumhuriyet Dans Gösterisi"),
            ("balo_davetiye.pdf", "pdf", "Cumhuriyet Balosu Davetiyesi"),
        ]),

        # --- COMBINED BATCH ROOT FOLDER (Yıllık_Toplu_Arşiv) ---
        ("Yıllık_Toplu_Arşiv/01.01.2024 Divan Toplantısı", [
            ("foto1.jpg", "image", "Foto 1"),
            ("rapor.pdf", "pdf", "2024 Raporu"),
        ]),
        ("Yıllık_Toplu_Arşiv/01.01.2025 Divan Toplantısı", [
            ("foto2.jpg", "image", "Foto 2"),
            ("rapor2025.pdf", "pdf", "2025 Raporu"),
        ]),
        ("Yıllık_Toplu_Arşiv/15.04.2024 Bahar Semineri", [
            ("seminer.jpg", "image", "Bahar Semineri 2024"),
            ("sunum.pdf", "pdf", "Bahar Semineri Sunumu"),
            ("notlar.docx", "docx", "Seminer Notları"),
        ]),
        ("Yıllık_Toplu_Arşiv/20.08.2024 Yaz Kampı", [
            ("kamp1.jpg", "image", "Yaz Kampı"),
            ("kamp_video.mp4", "video", "Kamp Video"),
        ]),
        ("Yıllık_Toplu_Arşiv/2024 Kış Paneli", [
            ("panel.jpg", "image", "Kış Paneli"),
        ]),
    ]

    for sub_dir, files in events_structure:
        full_sub_dir = os.path.join(abs_base, sub_dir)
        os.makedirs(full_sub_dir, exist_ok=True)

        for filename, ftype, label in files:
            file_path = os.path.join(full_sub_dir, filename)
            if ftype == "image":
                make_fake_image(file_path, label)
            elif ftype == "docx":
                make_fake_docx(file_path, label, f"{label} içeriği ve detayları.")
            elif ftype == "pdf":
                make_fake_pdf(file_path, label, f"{label} döküman metni.")
            elif ftype == "video":
                make_fake_mp4(file_path, 2)

    print("✅ Test arşivi başarıyla oluşturuldu!")
    print(f"📁 Konum: {abs_base}")
    print("Test klasörlerini görmek için:")
    print(f"  - {os.path.join(abs_base, '2024 Yılı Etkinlikleri')}")
    print(f"  - {os.path.join(abs_base, '2025 Yılı Etkinlikleri')}")
    print(f"  - {os.path.join(abs_base, 'Yıllık_Toplu_Arşiv')}")


if __name__ == "__main__":
    from PySide6 import QtGui
    app = QtGui.QGuiApplication.instance() or QtGui.QGuiApplication(sys.argv)
    create_test_archive()
